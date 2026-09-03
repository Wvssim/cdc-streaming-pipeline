from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "Assets"
OUTPUT = ROOT / "docs" / "Manuel_utilisation_plateforme_CDC.docx"
GITHUB = "https://github.com/Wvssim/cdc-streaming-pipeline"
OWNER = "Lazim Wassim"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cliquez ici puis mettez à jour le champ pour afficher le sommaire."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_link(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend((color, underline))
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_figure(doc, image, caption, width=6.35):
    path = ASSETS / image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_note(doc, title, text, color="EAF3F8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.add_run(title + " — ").bold = True
    p.add_run(text)


def add_steps(doc, steps):
    for number, text in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.1)
sec.right_margin = Cm(2.1)
sec.header_distance = Cm(0.8)
sec.footer_distance = Cm(0.8)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color in (("Title", 25, "123047"), ("Heading 1", 18, "123047"), ("Heading 2", 14, "007C83"), ("Heading 3", 12, "007C83")):
    style = styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(7)

styles["Caption"].font.name = "Times New Roman"
styles["Caption"].font.size = Pt(10)
styles["Caption"].font.italic = True
styles["Caption"].font.color.rgb = RGBColor(70, 70, 70)

doc.core_properties.title = "Manuel d’utilisation — Plateforme documentaire événementielle CDC"
doc.core_properties.author = OWNER
doc.core_properties.subject = "Guide utilisateur et d’exploitation de la plateforme CDC"
doc.core_properties.keywords = f"CDC, Kafka, microservices, EMSI, 6Solutions, {OWNER}, {GITHUB}"
doc.core_properties.comments = f"Travail original de {OWNER}. Dépôt officiel : {GITHUB}"

# En-tête et pied de page avec attribution persistante.
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = header.add_run("PLATEFORME DOCUMENTAIRE ÉVÉNEMENTIELLE — MANUEL D’UTILISATION")
hr.font.name = "Times New Roman"
hr.font.size = Pt(8)
hr.font.color.rgb = RGBColor(90, 90, 90)
footer_table = sec.footer.add_table(rows=1, cols=2, width=Cm(16.8))
footer_table.columns[0].width = Cm(13.5)
footer_table.columns[1].width = Cm(3.3)
left = footer_table.cell(0, 0).paragraphs[0]
left.add_run(f"© 2026 {OWNER} · {GITHUB}").font.size = Pt(8)
add_page_number(footer_table.cell(0, 1).paragraphs[0])

# Page de couverture.
logos = doc.add_table(rows=1, cols=2)
logos.alignment = WD_TABLE_ALIGNMENT.CENTER
logos.autofit = False
logos.columns[0].width = Cm(8)
logos.columns[1].width = Cm(8)
for cell in logos.rows[0].cells:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
logos.cell(0, 0).paragraphs[0].add_run().add_picture(str(ROOT / "docs" / "logo.png"), width=Inches(2.35))
logos.cell(0, 1).paragraphs[0].add_run().add_picture(str(ASSETS / "6solutions_logo.png"), width=Inches(2.25))
doc.add_paragraph("\n")
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("MANUEL D’UTILISATION")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plateforme documentaire événementielle fondée sur un pipeline CDC")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = RGBColor.from_string("007C83")
doc.add_paragraph("\n")
cover = doc.add_table(rows=5, cols=2)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.style = "Table Grid"
entries = (
    ("Auteur", OWNER),
    ("Filière", "DSI — 4ᵉ année, groupe G3"),
    ("Encadrement académique", "Mme Asmaa Roudane"),
    ("Encadrement professionnel", "M. Bekkali Mohamed"),
    ("Année universitaire", "2025–2026"),
)
for row, (label, value) in zip(cover.rows, entries):
    row.cells[0].paragraphs[0].add_run(label).bold = True
    row.cells[1].paragraphs[0].add_run(value)
    set_cell_shading(row.cells[0], "DDEBF0")
doc.add_paragraph("\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Dépôt officiel du projet :\n").bold = True
add_link(p, GITHUB, GITHUB)
p = doc.add_paragraph("Version 1.0 — septembre 2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True
doc.add_page_break()

doc.add_heading("Informations sur le document", level=1)
info = doc.add_table(rows=6, cols=2)
info.style = "Table Grid"
info.alignment = WD_TABLE_ALIGNMENT.CENTER
data = (
    ("Titre", "Manuel d’utilisation de la plateforme documentaire événementielle"),
    ("Version", "1.0"),
    ("Date", "3 septembre 2026"),
    ("Public concerné", "Utilisateur métier, encadrant, évaluateur et exploitant technique"),
    ("Auteur et propriétaire", OWNER),
    ("Référence officielle", GITHUB),
)
for row, (a, b) in zip(info.rows, data):
    row.cells[0].paragraphs[0].add_run(a).bold = True
    row.cells[1].paragraphs[0].add_run(b)
    set_cell_shading(row.cells[0], "EAF3F8")
doc.add_heading("Objet du manuel", level=2)
doc.add_paragraph("Ce manuel explique l’utilisation fonctionnelle de la plateforme, depuis l’authentification jusqu’au suivi des traitements asynchrones. Il présente également les interfaces techniques nécessaires à la démonstration et au diagnostic du pipeline de capture des changements.")
add_note(doc, "Attribution", f"Ce document et le projet associé constituent un travail original de {OWNER}. La provenance peut être vérifiée dans les métadonnées du document, son pied de page et le dépôt GitHub officiel.")

doc.add_heading("Table des matières", level=1)
add_toc(doc.add_paragraph())
doc.add_page_break()

doc.add_heading("1. Présentation générale", level=1)
doc.add_paragraph("La plateforme automatise les traitements déclenchés par le dépôt ou la modification d’un document. Une transaction enregistrée dans la base est capturée, publiée dans Kafka, puis consommée indépendamment par les services d’audit, de notification, d’intégrité, de reconnaissance optique de caractères et de détection d’incidents.")
doc.add_heading("1.1 Fonctions principales", level=2)
for item in ("Authentification sécurisée par jeton.", "Dépôt, consultation, téléchargement, renommage et suppression de documents.", "Extraction automatique du texte des documents compatibles.", "Traçabilité des opérations dans une piste d’audit.", "Vérification de l’intégrité par empreinte SHA-256 et chaîne de hachage.", "Notifications automatiques et consultation de leur statut.", "Détection et consultation des alertes de sécurité.", "Supervision des topics et groupes de consommateurs Kafka."):
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("1.2 Profils concernés", level=2)
doc.add_paragraph("L’utilisateur métier consulte et manipule les documents depuis l’interface Angular. L’exploitant technique utilise en complément Swagger, Kafbat UI, MailHog et MinIO afin de contrôler les API, les événements, les messages et le stockage.")
add_figure(doc, "screen_tableau_de_bord.png", "Figure 1 — Tableau de bord général de la plateforme", 6.4)
doc.add_page_break()

doc.add_heading("2. Accès et authentification", level=1)
doc.add_heading("2.1 Prérequis", level=2)
for item in ("Utiliser un navigateur récent, de préférence Microsoft Edge, Chrome ou Firefox.", "Vérifier que le frontend et les six services applicatifs sont démarrés.", "Vérifier que PostgreSQL, Kafka, le connecteur CDC, MinIO, MailHog et Kafbat UI sont disponibles."):
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("2.2 Connexion", level=2)
add_steps(doc, ("Ouvrir l’adresse http://localhost:4200.", "Saisir le nom d’utilisateur de démonstration : wassim.", "Saisir le mot de passe : wassim2026.", "Cliquer sur « Se connecter ».", "Vérifier l’affichage du tableau de bord."))
add_note(doc, "Sécurité", "Les identifiants fournis sont réservés à la démonstration locale. Ils ne doivent pas être réutilisés dans un environnement réel.", "FFF2CC")
add_figure(doc, "screen_tableau_de_bord.png", "Figure 2 — Écran obtenu après une authentification réussie", 6.35)
doc.add_page_break()

doc.add_heading("3. Gestion des documents", level=1)
doc.add_heading("3.1 Consulter les documents", level=2)
doc.add_paragraph("La rubrique « Documents » affiche les fichiers disponibles, leur taille, leur état, leur date de création et les principales actions. Les indicateurs du haut synthétisent le volume total et l’état des traitements.")
add_figure(doc, "screen_documents.png", "Figure 3 — Liste et indicateurs des documents", 5.5)
doc.add_heading("3.2 Déposer un document", level=2)
add_steps(doc, ("Ouvrir la rubrique « Documents ».", "Cliquer sur « Déposer un document ».", "Sélectionner un fichier autorisé depuis l’ordinateur.", "Confirmer le dépôt.", "Attendre le message de réussite, puis vérifier l’apparition du document dans la liste."))
add_note(doc, "Résultat attendu", "Le fichier est stocké dans MinIO, ses métadonnées sont enregistrées en base et les traitements asynchrones sont déclenchés automatiquement.")
doc.add_page_break()

doc.add_heading("3.3 Consulter la fiche détaillée", level=2)
doc.add_paragraph("Un clic sur un document ouvre sa fiche. Celle-ci rassemble ses informations générales, le texte extrait par reconnaissance optique et la chronologie des événements associés.")
add_figure(doc, "screen_detail.png", "Figure 4 — Fiche détaillée et résultat de l’extraction de texte", 5.8)
doc.add_heading("3.4 Télécharger, renommer ou supprimer", level=2)
add_steps(doc, ("Ouvrir la fiche du document concerné.", "Utiliser « Télécharger » pour récupérer le fichier original.", "Utiliser l’action de renommage pour modifier son nom logique, puis confirmer.", "Utiliser « Supprimer » uniquement si le retrait est souhaité, puis confirmer l’avertissement."))
add_note(doc, "Attention", "La suppression retire la ressource métier et génère un événement traçable. Vérifier l’identifiant et le nom du document avant confirmation.", "FCE4D6")
doc.add_page_break()

doc.add_heading("4. Suivi des traitements", level=1)
doc.add_heading("4.1 Piste d’audit", level=2)
doc.add_paragraph("La piste d’audit permet de contrôler les opérations réalisées et leur ordre d’arrivée. Elle facilite la démonstration de la traçabilité et le diagnostic d’une action utilisateur.")
add_figure(doc, "screen_piste_audit.png", "Figure 5 — Consultation de la piste d’audit", 5.7)
doc.add_heading("4.2 Notifications", level=2)
doc.add_paragraph("L’écran « Notifications » présente les messages générés par les événements documentaires. L’état affiché permet de confirmer que le consumer concerné a traité le message.")
add_figure(doc, "screen_notifications.png", "Figure 6 — Historique des notifications", 5.7)
doc.add_page_break()

doc.add_heading("4.3 Intégrité", level=2)
doc.add_paragraph("Le registre d’intégrité associe chaque document à une empreinte SHA-256 et l’inscrit dans une chaîne de hachage. Une chaîne déclarée intègre indique qu’aucune rupture n’a été détectée dans l’enchaînement des preuves.")
add_figure(doc, "screen_integrite.png", "Figure 7 — Registre d’intégrité et état de la chaîne de hachage", 5.7)
doc.add_heading("4.4 Alertes de sécurité", level=2)
doc.add_paragraph("La rubrique SIEM centralise les signaux suspects détectés par les règles de démonstration, par exemple un type de fichier interdit, une activité nocturne ou une succession rapide de dépôts.")
add_figure(doc, "screen_alertes_siem.png", "Figure 8 — Tableau des alertes de sécurité", 5.7)
doc.add_page_break()

doc.add_heading("5. Interfaces techniques de contrôle", level=1)
doc.add_heading("5.1 Documentation des API", level=2)
doc.add_paragraph("Swagger UI expose les routes de l’API documentaire. Cette interface permet de consulter les contrats, paramètres, formats de réponse et codes HTTP. L’exécution directe d’une route protégée nécessite un jeton valide.")
add_figure(doc, "swagger_ui_documents_api.png", "Figure 9 — Documentation interactive de l’API documentaire", 5.5)
doc.add_heading("5.2 Supervision du flux Kafka", level=2)
doc.add_paragraph("Kafbat UI permet de contrôler le cluster, les topics, les partitions et les groupes de consommateurs. Le topic source contient les changements capturés ; chaque consumer group représente la lecture indépendante d’un service.")
add_figure(doc, "kafbat_ui_docs_public_documents.png", "Figure 10 — Topic CDC contenant les événements documentaires", 6.2)
doc.add_page_break()

doc.add_heading("5.3 Groupes de consommateurs et rejeu", level=2)
doc.add_paragraph("La vue des consumer groups permet de vérifier les offsets et le retard éventuel. Un retard nul indique que le groupe a rejoint la dernière position disponible. Le rejeu consiste à repositionner un groupe afin qu’il retraite les événements, tout en conservant l’idempotence métier.")
add_figure(doc, "kafbat_consumer_groups_s4.png", "Figure 11 — Groupes de consommateurs indépendants", 6.25)
doc.add_heading("5.4 Stockage objet", level=2)
doc.add_paragraph("MinIO conserve le contenu binaire des documents dans le bucket « documents ». L’événement Kafka transporte la référence de l’objet plutôt que le fichier complet, conformément au modèle Claim Check.")
add_figure(doc, "Miniologed.png", "Figure 12 — Consultation du bucket documents dans MinIO", 6.25)
doc.add_page_break()

doc.add_heading("5.5 Contrôle des courriels", level=2)
doc.add_paragraph("MailHog intercepte les courriels de l’environnement de démonstration. Après un dépôt, la boîte de réception doit contenir une notification correspondant au document et à l’événement traité.")
add_figure(doc, "mailhog_inbox.png", "Figure 13 — Courriels de notification visibles dans MailHog", 6.25)

doc.add_heading("6. Scénario de démonstration recommandé", level=1)
add_steps(doc, ("Se connecter à l’interface Angular.", "Présenter les indicateurs du tableau de bord.", "Déposer un document de test.", "Ouvrir sa fiche, montrer le texte extrait puis télécharger l’original.", "Afficher successivement la piste d’audit, les notifications et le registre d’intégrité.", "Vérifier l’arrivée du courriel dans MailHog.", "Montrer l’événement et les consumer groups dans Kafbat UI.", "Afficher l’objet correspondant dans MinIO.", "Terminer par les alertes SIEM et le bilan des tests."))
add_note(doc, "Conseil", "Laisser deux à cinq secondes après le dépôt avant d’ouvrir les écrans de suivi, afin que tous les consumers terminent leur traitement.")
doc.add_page_break()

doc.add_heading("7. Résolution des problèmes fréquents", level=1)
issues = doc.add_table(rows=1, cols=3)
issues.style = "Table Grid"
issues.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, title in enumerate(("Symptôme", "Cause probable", "Action recommandée")):
    issues.cell(0, i).paragraphs[0].add_run(title).bold = True
    set_cell_shading(issues.cell(0, i), "DDEBF0")
rows = (
    ("Connexion refusée", "Identifiants incorrects ou API indisponible", "Utiliser wassim / wassim2026 et vérifier documents-api."),
    ("Dépôt impossible", "Format, taille ou service indisponible", "Essayer un petit PDF ou PNG et consulter le message affiché."),
    ("Document absent de la liste", "Transaction non terminée ou interface non actualisée", "Actualiser la page et vérifier PostgreSQL puis documents-api."),
    ("Texte OCR vide", "Document non compatible ou moteur OCR absent", "Tester une image nette et vérifier l’installation native de Tesseract."),
    ("Aucune notification", "Consumer ou MailHog indisponible", "Vérifier notification-service, son consumer group et MailHog."),
    ("Retard Kafka non nul", "Consumer arrêté ou erreur de traitement", "Redémarrer le service concerné et consulter le topic d’erreurs."),
    ("Fichier introuvable au téléchargement", "Objet MinIO absent", "Contrôler le bucket documents et la clé stockée en base."),
    ("Chaîne d’intégrité invalide", "Empreinte ou maillon incohérent", "Identifier le premier maillon fautif avant toute correction."),
)
for values in rows:
    cells = issues.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value

doc.add_heading("8. Bonnes pratiques", level=1)
for item in ("Ne jamais partager un jeton d’authentification dans une capture ou un dépôt public.", "Conserver les secrets dans des variables d’environnement.", "Utiliser des fichiers de test dépourvus de données personnelles.", "Vérifier les journaux avant de rejouer un consumer group.", "Sauvegarder PostgreSQL et MinIO avant toute opération destructive.", "Désactiver Swagger et les interfaces techniques dans un environnement exposé."):
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("9. Références et propriété", level=1)
doc.add_paragraph(f"Auteur : {OWNER}")
doc.add_paragraph("Établissement : EMSI Casablanca")
doc.add_paragraph("Organisme d’accueil : 6Solutions")
p = doc.add_paragraph("Dépôt GitHub officiel : ")
add_link(p, GITHUB, GITHUB)
doc.add_paragraph("Toute copie ou adaptation doit conserver une attribution explicite à l’auteur et une référence au dépôt officiel. Les noms et marques des technologies et organismes cités restent la propriété de leurs détenteurs respectifs.")
add_note(doc, "Fin du manuel", "Pour une démonstration visuelle complète, consulter également la vidéo Demonstration_Plateforme_CDC.mp4 fournie avec les livrables.")

doc.save(OUTPUT)
print(OUTPUT)
